import time
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from bs4 import BeautifulSoup
from urllib.parse import urljoin
from docx import Document
import pandas as pd
from flask import Flask, request, render_template, redirect, url_for
from contextlib import contextmanager

@contextmanager
def create_webdriver():
    driver = webdriver.Chrome()
    try:
        yield driver
    finally:
        driver.quit()


app = Flask(__name__)

@app.route("/", methods=["GET", "POST"])  # Allow both GET and POST requests
def index():
    if request.method == "POST":
        # Process the sign-in form data here
        # Assuming the sign-in is successful, you can redirect to the home page
        return redirect(url_for("home"))  # Redirect to the "home" route
    return render_template("index.html")  # Replace with your HTML filename


@app.route("/home", methods=["GET"])
def home():
    #if request.method == "POST":
    selected_course = request.args.get("selected_course")
    print(f"Selected Course URL: {selected_course}")
    if selected_course is not None:
        # Set pandas options to display the full link without truncation
        pd.set_option("display.max_colwidth", None)

        # base url
        base_url = "https://learn.codespace.co.za/login"
        
        with create_webdriver() as driver:
            # login
            driver.get(base_url)
            wait = WebDriverWait(driver, 10)

            username = wait.until(EC.presence_of_element_located((By.NAME, 'email')))
            password = wait.until(EC.presence_of_element_located((By.NAME, 'password')))

            username.send_keys('david@codespace.co.za')
            password.send_keys('Dave36code@#')

            login_button = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, '.button.is-primary')))
            login_button.click()

            time.sleep(15)  # increase sleep time

            # get all links function
            def get_all_links(url):
                driver.get(url)
                time.sleep(15)  # increase sleep time
                html = driver.page_source
                soup = BeautifulSoup(html, 'html.parser')
                anchors = soup.find_all('a')
                links = set()
                for anchor in anchors:
                    link = anchor.get('href')
                    if link and not link.startswith('javascript:;'):
                        full_link = urljoin(url, link)
                        links.add(full_link)
                return links
            
            time.sleep(30)  # increase sleep time for the web scraper to recognise the link parsed

            # list of course urls
            course_urls = [selected_course]

            link_info = []

            # Create a single Word document                                                                        
            doc = Document()

            # for each course url, get all links and test them
            for course_url in course_urls: 
                print(f"\nGetting links for: {course_url}")
                all_links = get_all_links(course_url)
                for link in all_links:
                    driver.get(link)
                    time.sleep(15)  # increase sleep time

                    # Check if page loaded properly
                    if "Page not found" in driver.title or "Error" in driver.title:
                        status = "Error"
                        text_count = video_count = 0
                        video_links = []
                    else:
                        status = "OK"
                        html = driver.page_source
                        soup = BeautifulSoup(html, 'html.parser')
                        content_blocks = soup.find_all(class_='lesson-content-block column')

                        text_count = sum(len(block.get_text().split()) for block in content_blocks)

                        video_elements = soup.find_all('iframe', src=lambda x: x and 'youtube.com' in x)
                        video_count = len(video_elements)
                        video_links = [el['src'] for el in video_elements]

                        h4_elements = soup.find_all('h4')
                        h4_titles = [h4.get_text().strip() for h4 in h4_elements]

                        main_text = [block.get_text().strip() for block in content_blocks]

                        # Add titles and text to the Word document
                        for title, text in zip(h4_titles, main_text):
                            doc.add_heading(title, level=1)
                            doc.add_paragraph(text)

                    link_info.append({
                        "Link": link,
                        "Status": status,
                        "Text Count": text_count,
                        "Video Count": video_count,
                        "Video Links": video_links,
                        "H4": h4_titles,
                        "main_text": main_text,
                    })

            # Print summary
            print(f"\nTotal links: {len(all_links)}")
            
            # Set the pandas display option to show all columns and rows
            pd.set_option("display.max_columns", None)
            pd.set_option("display.max_rows", None)

            # Create DataFrame with link info
            df = pd.DataFrame(link_info, columns=["Link", "Status", "Text Count", "Video Count", "Video Links", "H4", "main_text"])
            print(df)
            
            # Save the single Word document
            doc.save("titles_and_text3.docx")
            
            # Calculate summary statistics
            total_links = len(all_links)
            total_text_count = df["Text Count"].sum()
            total_video_count = df["Video Count"].sum()
            total_broken_links = df[df["Status"] == "Error"].shape[0]
            # total_image_count = ...  # You need to calculate this based on your data
            
            # close WebDriver
            driver.quit()
            
            # Pass data to the template
            return render_template(
                "home_page.html",
                total_links=total_links,
                total_text_count=total_text_count,
                total_video_count=total_video_count,
                total_broken_links=total_broken_links,
                # total_image_count=total_image_count,
                link_infos=link_info
            )
        return render_template("home_page.html")  # Replace with your HTML filename
    return render_template("home_page.html")  # Replace with your HTML filename
if __name__ == "__main__":
    app.run(debug=True)

